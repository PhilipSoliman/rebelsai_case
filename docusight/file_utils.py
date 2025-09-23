import csv
import os
import shutil
import uuid
import zipfile
from pathlib import Path
from typing import Optional

import aiofiles
import PyPDF2
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from dropbox import Dropbox, files
from fastapi import HTTPException, UploadFile
from fastapi.concurrency import run_in_threadpool
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from striprtf.striprtf import rtf_to_text

from docusight.config import settings
from docusight.logging import logger
from docusight.models import Document, Folder, User


class MetaDict(dict):
    filename: str
    path: str
    size: int
    created: float
    modified: float

    def __init__(
        self, filename: str, path: str, size: int, created: float, modified: float
    ):
        super().__init__(
            filename=filename, path=path, size=size, created=created, modified=modified
        )

    def __getattr__(self, item):
        return self[item]


async def add_zipped_folder_to_database(
    zipped_folder: UploadFile,
    db: AsyncSession,
    dropbox_client: Dropbox,
    user: User,
    drill: bool,
) -> Folder:
    """
    Extracts a zipped folder, adds its contents to the database, uploads files to Dropbox, and cleans up temporary files.

    Args:
        request (Request): FastAPI request object.
        zipped_folder (UploadFile): The uploaded zipped folder (.zip file).
        db (AsyncSession): Database session.
        drill (bool): Whether to recursively add subfolders.

    Returns:
        Folder: The Folder ORM instance added to the database.
    """
    # construct temp paths
    tmp_dir = generate_tmp_dir(user)
    os.makedirs(tmp_dir, exist_ok=True)
    try:
        tmp_zipped_folder_path = Path(tmp_dir) / zipped_folder.filename
        tmp_client_dir = tmp_zipped_folder_path.with_suffix(
            ""
        )  # folder name without .zip

        # extract zip to temp directory
        await write_zip_in_chunks(zipped_folder, tmp_zipped_folder_path)
        await run_in_threadpool(extract_zip, tmp_zipped_folder_path)
        await delete_zipfile(tmp_zipped_folder_path)

        # convert files to plain text (+ overwrite paths)
        tmp_client_paths = [Path(p) for p in tmp_client_dir.rglob("*.*")]
        tmp_client_txt_file_map = await convert_files_to_plain_text(
            tmp_client_paths, tmp_dir
        )

        # add folder to database
        folder = await add_folder_to_database(
            current_dir=tmp_client_dir,
            tmp_dir=tmp_dir,
            db=db,
            user_id=user.id,
            drill=drill,
            original_file_map=tmp_client_txt_file_map,
        )

        # upload plain text paths to dropbox
        tmp_client_txt_files = list(tmp_client_txt_file_map.keys())
        dropbox_paths = await upload_files_to_dropbox(
            dropbox_client, tmp_client_txt_files, tmp_dir
        )
        for relative_path, dropbox_path in dropbox_paths.items():
            local_txt_path = tmp_dir / relative_path
            original_meta = tmp_client_txt_file_map.get(local_txt_path)
            document = await get_document_by_path(str(original_meta.path), db, user)
            if document:
                document.dropbox_path = dropbox_path
                db.add(document)

        # remove temp client directory
        await run_in_threadpool(shutil.rmtree, tmp_client_dir)

    except Exception as e:
        logger.error(f"Error adding zipped folder to database: {e}")
        await run_in_threadpool(shutil.rmtree, tmp_client_dir)
        raise HTTPException(
            status_code=500, detail="Failed to process and add zipped folder"
        )
    return folder


def generate_tmp_dir(user: User) -> Path:
    """
    Generates a temporary directory path for a user based on their ID and display name.
    Args:
        user (User): The User ORM instance.
    Returns:
        Path: The generated temporary directory path.
    """
    user_name: str = user.display_name
    user_name_fmt = user_name.lower().replace(" ", "_")
    return Path(settings.TEMP_DIR) / f"{user.id}_{user_name_fmt}"


async def delete_zipfile(zip_path: Path):
    """
    Asynchronously deletes a zip file from disk.

    Args:
        zip_path (Path): Path to the zip file to delete.
    """
    if zip_path.exists():
        await run_in_threadpool(zip_path.unlink)


async def write_zip_in_chunks(upload_file: UploadFile, dest_path: Path):
    """
    Asynchronously writes an uploaded zip file to disk in chunks.

    Args:
        upload_file (UploadFile): The uploaded zip file.
        dest_path (Path): Destination path to write the file.
    """
    with open(dest_path, "wb") as f:
        while True:
            chunk = await upload_file.read(settings.ZIP_FILE_READ_CHUNK_SIZE)
            if not chunk:
                break
            f.write(chunk)


def extract_zip(tmp_zipped_folder_path: Path):
    """
    Extracts a zip file to a directory with the same name (without .zip).

    Args:
        tmp_zipped_folder_path (Path): Path to the zip file to extract.
    """
    with zipfile.ZipFile(tmp_zipped_folder_path, "r") as zip_ref:
        zip_ref.extractall(tmp_zipped_folder_path.parent)


async def convert_files_to_plain_text(
    paths: list[Path], tmp_dir: Path
) -> dict[Path, MetaDict]:
    """
    Converts multiple files to plain text if supported. Deletes old files and saves plain text versions.

    Supported types: .txt, .docx, .pdf, .csv, .rtf, .html

    Args:
        paths (list[Path]): List of file paths to convert.
    Returns:
        dict[Path, dict]: Mapping of plain text file paths to original file metadata.
    """
    plain_text_files = {}
    for path in paths:
        text = await file_to_plain_text(path)

        if text is not None:
            original_stats = path.stat()
            original_meta = MetaDict(
                filename=path.name,
                path=str(path.relative_to(tmp_dir)),
                size=original_stats.st_size,
                created=original_stats.st_ctime,
                modified=original_stats.st_mtime,
            )
            plain_text_path = path.with_suffix(".txt")
            plain_text_files[plain_text_path] = original_meta
            async with aiofiles.open(
                plain_text_path, "w", encoding="utf-8", errors="replace"
            ) as f:
                await f.write(text)
            await run_in_threadpool(path.unlink)
    return plain_text_files


async def file_to_plain_text(path: Path) -> Optional[str]:
    """
    Converts a file to plain text if supported. Returns None if not supported or conversion fails.

    Supported types: .txt, .docx, .pdf, .csv, .rtf, .html
    """
    ext = path.suffix.lower()
    try:
        if ext == ".txt":
            async with aiofiles.open(
                path, "r", encoding="utf-8", errors="replace"
            ) as f:
                return await f.read()
        elif ext == ".docx":
            return await run_in_threadpool(read_docx, path)
        elif ext == ".pdf":
            return await run_in_threadpool(read_pdf, path)
        elif ext == ".csv":
            async with aiofiles.open(
                path, "r", encoding="utf-8", errors="replace"
            ) as f:
                content = await f.read()
            return await run_in_threadpool(parse_csv, content)
        elif ext in [".htm", ".html"]:
            async with aiofiles.open(
                path, "r", encoding="utf-8", errors="replace"
            ) as f:
                html = await f.read()
            return await run_in_threadpool(parse_html, html)
        elif ext == ".rtf":
            async with aiofiles.open(
                path, "r", encoding="utf-8", errors="replace"
            ) as f:
                rtf = await f.read()
            return await run_in_threadpool(rtf_to_text, rtf)
        else:
            # unsupported type
            return None
    except Exception as e:
        logger.error(f"Error converting file {path}: {e}")
        return None


def read_docx(p):
    doc = DocxDocument(str(p))
    return "\n".join([para.text for para in doc.paragraphs])


def read_pdf(p):
    text = ""
    with open(p, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text


def parse_csv(data):
    lines = []
    for row in csv.reader(data.splitlines()):
        lines.append(",".join(row))
    return "\n".join(lines)


def parse_html(data):
    soup = BeautifulSoup(data, "html.parser")
    return soup.get_text()


async def add_folder_to_database(
    current_dir: Path,
    tmp_dir: Path,
    db: AsyncSession,
    user_id: int,
    drill: bool,
    original_file_map: dict[Path, MetaDict],
    parent_folder: Optional[Folder] = None,
) -> Folder:
    """
    Recursively adds a folder and its documents to the database.

    Args:
        current_dir (Path): Path to the current folder.
        tmp_dir (Path): Base temp directory.
        db (AsyncSession): Database session.
        drill (bool): Whether to recursively add subfolders.
        file_paths (list[Path]): List to collect file paths for upload.
        parent_folder (Optional[Folder]): Parent folder ORM instance.

    Returns:
        Folder: The Folder ORM instance added to the database.
    """
    folder = Folder(
        path=str(current_dir.relative_to(tmp_dir)),
        name=current_dir.name,
        parent_id=parent_folder.id if parent_folder else None,
        user_id=user_id,
    )
    db.add(folder)
    await db.flush()
    for file in os.listdir(current_dir):
        file_path = Path(current_dir) / file
        if file_path.is_file():
            original_meta = original_file_map[file_path]
            document = Document(
                folder=folder,
                **original_meta,
                dropbox_path=None,  # NOTE: to be updated after upload
                plain_text_size=file_path.stat().st_size,
                user_id=user_id,
            )
            db.add(document)

    # Add subfolders recursively
    if drill:
        for subfolder_name in os.listdir(current_dir):
            subfolder_path = Path(current_dir) / subfolder_name
            if subfolder_path.is_dir():
                _ = await add_folder_to_database(
                    current_dir=subfolder_path,
                    tmp_dir=tmp_dir,
                    db=db,
                    user_id=user_id,
                    drill=drill,
                    original_file_map=original_file_map,
                    parent_folder=folder,
                )
    return folder


async def upload_files_to_dropbox(
    dropbox_client: Dropbox, file_paths: list[Path], tmp_dir: Path
) -> tuple[dict, files.UploadSessionFinishBatchResult]:
    """
    Uploads multiple files to Dropbox using upload sessions, chunking each file and finishing all sessions in a batch.

    Args:
        request (Request): FastAPI request object.
        file_paths (list[Path]): List of file paths to upload.
        tmp_dir (Path): Base temp directory.

    Returns:
        tuple[dict, files.UploadSessionFinishBatchResult]:
            - dict mapping relative file paths to Dropbox paths
            - Dropbox batch result object
    """
    try:
        dropbox_paths = {}
        entries: list[files.UploadSessionFinishArg] = []
        batch_start_result = await run_in_threadpool(
            dropbox_client.files_upload_session_start_batch,
            len(file_paths),
            files.UploadSessionType.sequential,
        )

        # get session ids
        if len(batch_start_result.session_ids) != len(file_paths):
            raise HTTPException(
                status_code=500,
                detail="Dropbox did not return session IDs for all files",
            )

        # Loop over files and upload each one in chunks
        for file_path, session_id in zip(file_paths, batch_start_result.session_ids):
            # Relative filepath
            relative_path = file_path.relative_to(tmp_dir)

            # Generate unique file ID and path
            file_id = str(uuid.uuid4())
            file_ext = (
                "." + file_path.name.split(".")[-1] if "." in file_path.name else ""
            )
            dropbox_filename = f"{file_id}{file_ext}"
            dropbox_path = f"{settings.UPLOAD_DIR}/{dropbox_filename}"

            # Append the rest of the chunks using aiofiles for async file I/O
            cursor = files.UploadSessionCursor(session_id=session_id, offset=0)
            async with aiofiles.open(file_path, "rb") as file:
                content = await file.read()
                await run_in_threadpool(
                    dropbox_client.files_upload_session_append_v2,
                    content,
                    cursor,
                )
                cursor.offset += len(content)
            await run_in_threadpool(
                dropbox_client.files_upload_session_append_v2,
                b"",
                cursor,
                close=True,
            )

            # Save entry for batch finish
            commit = files.CommitInfo(path=dropbox_path)
            finish_arg = files.UploadSessionFinishArg(cursor=cursor, commit=commit)
            entries.append(finish_arg)

            # Save metadata for response
            dropbox_paths[relative_path] = dropbox_path

        # Finish all sessions in one batch
        finish_batch_result: files.UploadSessionFinishBatchResult = (
            await run_in_threadpool(
                dropbox_client.files_upload_session_finish_batch_v2,
                entries,
            )
        )

    except Exception as e:
        logger.error(f"Error uploading file(s) to Dropbox: {e}")
        raise HTTPException(
            status_code=500, detail="Failed to upload file(s) to Dropbox"
        )

    # check for any failed uploads
    failures = []
    for entry in finish_batch_result.entries:
        if entry.is_failure():
            logger.error(f"Failed to upload file to Dropbox: {entry}")
            failures.append(entry)
    if failures:
        raise HTTPException(
            status_code=500, detail="Failed to upload one or more files to Dropbox"
        )

    return dropbox_paths


async def get_folder_by_path(
    path: str, db: AsyncSession, user: User
) -> Optional[Folder]:
    """
    Retrieves a Folder ORM instance by its path from the database.

    Args:
        path (str): Relative path of the folder.
        db (AsyncSession): Database session.

    Returns:
        Optional[Folder]: The Folder ORM instance, or None if not found.
    """
    segments = Path(path).parts
    return await get_folder_by_segments(segments, db, user)


async def get_document_by_path(
    path: str, db: AsyncSession, user: User
) -> Optional[Document]:
    """
    Retrieves a Document ORM instance by its path from the database.

    Args:
        path (str): Relative path of the document.
        db (AsyncSession): Database session.

    Returns:
        Optional[Document]: The Document ORM instance, or None if not found.
    """
    segments = Path(path).parts
    if len(segments) < 2:
        raise HTTPException(status_code=400, detail="Invalid document path")

    # get parent folder of document
    *folder_segments, filename = segments
    folder = await get_folder_by_segments(folder_segments, db, user)
    if not folder:
        raise HTTPException(status_code=404, detail="Document folder not found")

    # get document by filename and folder_id
    result = await db.execute(
        select(Document).where(
            Document.folder_id == folder.id, Document.filename == filename
        )
    )
    return result.scalars().first()


async def get_folder_by_segments(
    segments: list[str], db: AsyncSession, user: User
) -> Optional[Folder]:
    """
    Retrieves a Folder ORM instance by its path segments from the database.

    Args:
        segments (list[str]): List of path segments representing the folder hierarchy.
        db (AsyncSession): Database session.

    Returns:
        Optional[Folder]: The Folder ORM instance, or None if not found.
    """
    parent_id = None
    folder = None
    for segment in segments:
        # only search parent folder
        query = select(Folder).where(
            Folder.parent_id == parent_id, Folder.user_id == user.id
        )

        # search for the folder by name
        query = query.where(Folder.name == segment)
        result = await db.execute(query)
        folder = result.scalars().first()
        if not folder:
            return None
        parent_id = folder.id
    return folder


async def get_documents_in_folder(
    folder: Folder, db: AsyncSession, drill: bool = False, classified: bool = False
) -> list[Document]:
    """
    Retrieves all Document ORM instances in a given folder.

    Args:
        folder (Folder): The Folder ORM instance.
        db (AsyncSession): Database session.

    Returns:
        list[Document]: List of Document ORM instances in the folder.
    """
    documents = []
    query = select(Document).where(
        Document.folder_id == folder.id,
        (
            Document.classification != None
            if classified
            else Document.classification == None
        ),
    )
    if drill:
        # Get documents in the current folder
        result = await db.execute(query)

        # Recursively get documents in subfolders
        subfolders = await get_subfolders_in_folder(folder, db)
        for subfolder in subfolders:
            documents.extend(
                await get_documents_in_folder(
                    subfolder, db, drill=True, classified=classified
                )
            )
    else:
        result = await db.execute(query)

    documents.extend(result.scalars().all())

    return documents


async def get_subfolders_in_folder(folder: Folder, db: AsyncSession) -> list[Folder]:
    """
    Retrieves all subfolder ORM instances for a given folder.

    Args:
        folder (Folder): The Folder ORM instance.
        db (AsyncSession): Database session.

    Returns:
        list[Folder]: List of subfolder ORM instances.
    """
    result = await db.execute(select(Folder).where(Folder.parent_id == folder.id))
    return result.scalars().all()


async def download_files_from_dropbox(
    dropbox_client: Dropbox, dropbox_paths: list[str], tmp_dir: Path
) -> list[Path]:
    """
    Downloads multiple files from Dropbox and returns their contents as a list of strings.

    Args:
        dropbox_client (Dropbox): Authenticated Dropbox client.
        dropbox_paths (list[str]): List of Dropbox file paths.
        tmp_dir (Path): Temporary directory to save downloaded files.

    Returns:
        list[str]: List of file contents as strings.
    """
    download_paths = []
    for dropbox_path in dropbox_paths:
        download_path = str(tmp_dir / Path(dropbox_path).name)
        download_paths.append(download_path)
        await run_in_threadpool(
            dropbox_client.files_download_to_file, download_path, dropbox_path
        )

    return download_paths
